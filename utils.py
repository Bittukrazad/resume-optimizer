import re
import io
import logging
from rapidfuzz import fuzz, process
from typing import Dict, List, Tuple, Optional

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# ============================================================
# ENHANCED SECTION PARSER (Multi-Column + Dynamic Detection)
# ============================================================

# Comprehensive section headers with variations
SECTION_HEADERS = {
    "summary": [
        "summary", "professional summary", "career summary", "objective",
        "profile", "about me", "career objective", "professional profile",
        "executive summary", "introduction", "bio", "overview"
    ],
    "skills": [
        "skills", "technical skills", "skills & tools", "core competencies",
        "technologies", "expertise", "proficiencies", "technical proficiencies",
        "key skills", "skillset", "technical expertise", "tools", "programming skills",
        "languages", "frameworks", "technical competencies"
    ],
    "experience": [
        "experience", "work experience", "professional experience", "employment history",
        "work history", "career history", "employment", "professional background",
        "relevant experience", "internships", "internship experience"
    ],
    "projects": [
        "projects", "academic projects", "personal projects", "key projects",
        "portfolio", "project work", "major projects", "research projects",
        "capstone projects", "selected projects"
    ],
    "education": [
        "education", "academic background", "qualifications", "educational background",
        "academic qualifications", "degrees", "schooling", "academics"
    ],
    "certifications": [
        "certifications", "courses", "licenses", "training", "professional development",
        "certificates", "credentials", "licensed", "accreditations", "online courses"
    ],
    "achievements": [
        "achievements", "awards", "accomplishments", "honors", "recognition",
        "honors & awards", "distinctions", "accolades", "achievements & awards"
    ],
    "languages": [
        "languages", "language proficiency", "linguistic skills", "language skills",
        "spoken languages", "languages known"
    ],
    "publications": [
        "publications", "research papers", "papers", "published work",
        "research", "articles", "journal publications"
    ],
    "volunteering": [
        "volunteering", "volunteer work", "volunteer experience", "community service",
        "social work", "extracurricular activities", "activities"
    ],
    "interests": [
        "interests", "hobbies", "personal interests", "hobbies & interests",
        "extracurricular", "activities & interests"
    ],
    "references": [
        "references", "referees", "professional references"
    ]
}

def find_best_section(header_text: str) -> Optional[str]:
    """Enhanced fuzzy matching with multi-word support"""
    header_text = header_text.lower().strip()
    
    # Remove common punctuation that might interfere
    header_text = re.sub(r'[:\-–—_|]', ' ', header_text)
    header_text = ' '.join(header_text.split())
    
    best_match = None
    best_score = 0

    for canonical, variations in SECTION_HEADERS.items():
        for v in variations:
            # Try exact match first
            if header_text == v:
                return canonical
            
            # Fuzzy match
            score = fuzz.ratio(header_text, v)
            
            # Partial match for multi-word headers
            partial_score = fuzz.partial_ratio(header_text, v)
            
            # Token sort for word order independence
            token_score = fuzz.token_sort_ratio(header_text, v)
            
            # Take best score
            max_score = max(score, partial_score, token_score)
            
            if max_score > best_score:
                best_match = canonical
                best_score = max_score

    # More lenient threshold for better detection
    return best_match if best_score >= 55 else None


def detect_column_layout(text: str) -> bool:
    """Detect if resume has two-column layout"""
    lines = text.split('\n')
    
    # Check for multiple short lines with content on both sides
    multi_column_indicators = 0
    
    for line in lines[:50]:  # Check first 50 lines
        # Look for lines with content, then spaces, then more content
        if re.search(r'\w+\s{5,}\w+', line):
            multi_column_indicators += 1
    
    return multi_column_indicators > 3


def split_two_column_text(text: str) -> str:
    """Split two-column layout into single column"""
    lines = text.split('\n')
    processed_lines = []
    
    for line in lines:
        # Find positions where there are 5+ consecutive spaces
        parts = re.split(r'\s{5,}', line)
        
        # Add each part as separate line
        for part in parts:
            if part.strip():
                processed_lines.append(part.strip())
    
    return '\n'.join(processed_lines)


def clean_text(t: str) -> str:
    """Clean text while preserving structure"""
    if not t:
        return ""
    
    # Normalize whitespace but keep newlines
    t = re.sub(r' +', ' ', t)
    t = re.sub(r'\n\s*\n\s*\n+', '\n\n', t)
    
    return t.strip()


def parse_resume_sections(text: str) -> Dict[str, str]:
    """
    Universal parser that works for:
    - Two-column resumes
    - Mixed-case section headers
    - Headers with or without ':' or '|' or '-'
    - Headers on left/right columns
    - Overleaf & Canva templates
    - ATS-friendly single column
    - Creative layouts
    """
    
    # Store original text
    original_text = text
    
    # Normalize text
    text = text.replace("\r", "\n")
    
    # Detect and handle two-column layout
    if detect_column_layout(text):
        logger.info("📊 Two-column layout detected, converting to single column")
        text = split_two_column_text(text)
    
    # Convert ALL CAPS headings to Title Case
    text = re.sub(
        r"\n([A-Z][A-Z ]{3,})\n",
        lambda m: "\n" + m.group(1).title() + "\n",
        text
    )
    
    # Enhanced heading pattern - detects multiple formats
    heading_patterns = [
        # Standard newline-separated headers
        r"\n\s*([A-Za-z][A-Za-z &/]{2,40})\s*\n",
        
        # Headers with colon
        r"\n\s*([A-Za-z][A-Za-z &/]{2,40})\s*:",
        
        # Headers with pipe separator
        r"\n\s*([A-Za-z][A-Za-z &/]{2,40})\s*\|",
        
        # Headers with dash/underscore
        r"\n\s*([A-Za-z][A-Za-z &/]{2,40})\s*[-–—_]{2,}",
        
        # All caps headers (before conversion)
        r"\n\s*([A-Z][A-Z ]{3,40})\s*\n",
        
        # Headers at start of line with colon
        r"^([A-Za-z][A-Za-z &/]{2,40})\s*:",
    ]
    
    # Find all potential headers using all patterns
    all_matches = []
    
    for pattern in heading_patterns:
        matches = list(re.finditer(pattern, text, re.MULTILINE))
        for match in matches:
            header_text = match.group(1).strip()
            
            # Filter out common false positives
            if len(header_text.split()) <= 5:  # Max 5 words for header
                all_matches.append({
                    'text': header_text,
                    'start': match.start(),
                    'end': match.end()
                })
    
    # Sort by position
    all_matches.sort(key=lambda x: x['start'])
    
    # Remove duplicates (same position)
    unique_matches = []
    last_pos = -100
    
    for match in all_matches:
        if match['start'] - last_pos > 5:  # At least 5 chars apart
            unique_matches.append(match)
            last_pos = match['start']
    
    # Initialize sections
    sections = {k: "" for k in SECTION_HEADERS.keys()}
    
    if not unique_matches:
        logger.warning("⚠️ No section headers detected")
        # Try to extract at least some content
        sections["summary"] = text[:500] if len(text) > 500 else text
        return sections
    
    # Process each match
    for i, match in enumerate(unique_matches):
        header_raw = match['text']
        
        if not header_raw:
            continue
        
        # Find content between this header and next
        start = match['end']
        end = unique_matches[i+1]['start'] if i+1 < len(unique_matches) else len(text)
        
        body = text[start:end].strip()
        
        # Match to canonical section
        best = find_best_section(header_raw)
        
        if best:
            # Append content (allow multiple sections with same name)
            if sections[best]:
                sections[best] += "\n\n" + clean_text(body)
            else:
                sections[best] = clean_text(body)
            
            logger.info(f"✅ Detected: {header_raw} → {best} ({len(body)} chars)")
        else:
            logger.debug(f"❌ Unmatched header: {header_raw}")
    
    # Post-processing: If sections are empty, try alternate extraction
    if not any(sections.values()):
        logger.warning("⚠️ No content extracted, trying alternate method")
        sections = fallback_extraction(original_text)
    
    return sections


def fallback_extraction(text: str) -> Dict[str, str]:
    """Fallback method for difficult resumes"""
    sections = {k: "" for k in SECTION_HEADERS.keys()}
    
    # Split by double newlines
    blocks = text.split('\n\n')
    
    for block in blocks:
        block = block.strip()
        if not block or len(block) < 20:
            continue
        
        # Check first line for section name
        first_line = block.split('\n')[0].strip()
        
        matched_section = find_best_section(first_line)
        
        if matched_section:
            content = '\n'.join(block.split('\n')[1:])
            sections[matched_section] += "\n" + clean_text(content)
        else:
            # Add to summary if no match
            sections["summary"] += "\n" + clean_text(block)
    
    return sections


# ============================================================
# ROBUST FILE EXTRACTION
# ============================================================

def extract_text_from_pdf(file) -> str:
    """Robust PDF extraction with multiple fallback methods"""
    text = ""
    methods_tried = []
    
    # Method 1: PyPDF2
    try:
        from PyPDF2 import PdfReader
        file.seek(0)
        pdf_reader = PdfReader(file)
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        
        if text.strip():
            logger.info("✅ PyPDF2 extraction successful")
            return clean_extracted_text(text)
        methods_tried.append("PyPDF2 (no text)")
    except Exception as e:
        methods_tried.append(f"PyPDF2 (error: {str(e)[:50]})")
        logger.warning(f"⚠️ PyPDF2 failed: {str(e)[:100]}")
    
    # Method 2: pdfplumber
    try:
        import pdfplumber
        file.seek(0)
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        
        if text.strip():
            logger.info("✅ pdfplumber extraction successful")
            return clean_extracted_text(text)
        methods_tried.append("pdfplumber (no text)")
    except ImportError:
        methods_tried.append("pdfplumber (not installed)")
    except Exception as e:
        methods_tried.append(f"pdfplumber (error: {str(e)[:50]})")
        logger.warning(f"⚠️ pdfplumber failed: {str(e)[:100]}")
    
    # Method 3: pypdf
    try:
        from pypdf import PdfReader as PyPdfReader
        file.seek(0)
        pdf_reader = PyPdfReader(file)
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        
        if text.strip():
            logger.info("✅ pypdf extraction successful")
            return clean_extracted_text(text)
        methods_tried.append("pypdf (no text)")
    except ImportError:
        methods_tried.append("pypdf (not installed)")
    except Exception as e:
        methods_tried.append(f"pypdf (error: {str(e)[:50]})")
        logger.warning(f"⚠️ pypdf failed: {str(e)[:100]}")
    
    # Method 4: PDFMiner
    try:
        from pdfminer.high_level import extract_text as pdfminer_extract
        file.seek(0)
        text = pdfminer_extract(file)
        
        if text.strip():
            logger.info("✅ pdfminer extraction successful")
            return clean_extracted_text(text)
        methods_tried.append("pdfminer (no text)")
    except ImportError:
        methods_tried.append("pdfminer (not installed)")
    except Exception as e:
        methods_tried.append(f"pdfminer (error: {str(e)[:50]})")
        logger.warning(f"⚠️ pdfminer failed: {str(e)[:100]}")
    
    # All methods failed
    error_msg = f"Unable to extract text from PDF.\n\nMethods tried:\n" + "\n".join(f"• {m}" for m in methods_tried)
    error_msg += "\n\nPossible reasons:\n• PDF is password-protected\n• PDF contains only images (scanned)\n• PDF is corrupted"
    
    logger.error("❌ All PDF extraction methods failed")
    raise ValueError(error_msg)


def extract_text_from_docx(file) -> str:
    """Robust DOCX extraction with fallback methods"""
    text = ""
    methods_tried = []
    
    # Method 1: python-docx
    try:
        from docx import Document
        file.seek(0)
        doc = Document(file)
        
        # Extract from paragraphs
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        text = "\n".join(paragraphs)
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text += "\n" + "     ".join(row_text)  # 5 spaces for column separation
        
        if text.strip():
            logger.info("✅ python-docx extraction successful")
            return clean_extracted_text(text)
        methods_tried.append("python-docx (no text)")
    except Exception as e:
        methods_tried.append(f"python-docx (error: {str(e)[:50]})")
        logger.warning(f"⚠️ python-docx failed: {str(e)[:100]}")
    
    # Method 2: docx2txt
    try:
        import docx2txt
        file.seek(0)
        text = docx2txt.process(file)
        
        if text.strip():
            logger.info("✅ docx2txt extraction successful")
            return clean_extracted_text(text)
        methods_tried.append("docx2txt (no text)")
    except ImportError:
        methods_tried.append("docx2txt (not installed)")
    except Exception as e:
        methods_tried.append(f"docx2txt (error: {str(e)[:50]})")
        logger.warning(f"⚠️ docx2txt failed: {str(e)[:100]}")
    
    # Method 3: XML extraction
    try:
        import zipfile
        from xml.etree import ElementTree as ET
        
        file.seek(0)
        with zipfile.ZipFile(file, 'r') as docx_zip:
            xml_content = docx_zip.read('word/document.xml')
            tree = ET.fromstring(xml_content)
            
            namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            paragraphs = tree.findall('.//w:p', namespaces)
            
            for para in paragraphs:
                texts = para.findall('.//w:t', namespaces)
                para_text = ''.join([t.text for t in texts if t.text])
                if para_text.strip():
                    text += para_text + "\n"
        
        if text.strip():
            logger.info("✅ XML extraction successful")
            return clean_extracted_text(text)
        methods_tried.append("XML extraction (no text)")
    except Exception as e:
        methods_tried.append(f"XML extraction (error: {str(e)[:50]})")
        logger.warning(f"⚠️ XML extraction failed: {str(e)[:100]}")
    
    # All methods failed
    error_msg = f"Unable to extract text from DOCX.\n\nMethods tried:\n" + "\n".join(f"• {m}" for m in methods_tried)
    error_msg += "\n\nPossible reasons:\n• File is corrupted\n• File is password-protected\n• Not a valid DOCX format"
    
    logger.error("❌ All DOCX extraction methods failed")
    raise ValueError(error_msg)


def clean_extracted_text(text: str) -> str:
    """Clean extracted text while preserving structure"""
    if not text:
        return ""
    
    # Replace URLs with placeholder
    text = re.sub(r'http[s]?://\S+', ' [URL] ', text)
    
    # Replace email with placeholder
    text = re.sub(r'\S+@\S+\.\S+', ' [EMAIL] ', text)
    
    # Remove problematic special characters but keep structure
    text = re.sub(r'[^\w\s\.\,\-\(\)\[\]\+\#\:\;\/\&\%\@\*\'\"\n]', ' ', text)
    
    # Normalize whitespace
    text = re.sub(r' +', ' ', text)
    text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
    
    return text.strip()


def validate_resume_content(text: str) -> Tuple[bool, str]:
    """Validate if extracted text is actually a resume"""
    if not text or len(text.strip()) < 50:
        return False, "Extracted text is too short (less than 50 characters)."
    
    # Comprehensive resume indicators
    resume_indicators = [
        # Section headers
        'experience', 'education', 'skills', 'project', 'summary', 
        'objective', 'profile', 'qualification', 'certification',
        'achievement', 'internship', 'training', 'award',
        
        # Educational terms
        'bachelor', 'master', 'degree', 'university', 'college',
        'school', 'graduate', 'undergraduate', 'btech', 'mtech',
        'bsc', 'msc', 'diploma', 'cgpa', 'gpa', 'percentage',
        
        # Work terms
        'work', 'intern', 'job', 'role', 'position', 'responsibilities',
        'company', 'organization', 'team', 'developed', 'managed',
        'led', 'designed', 'implemented', 'built', 'created',
        
        # Technical terms
        'python', 'java', 'javascript', 'react', 'node', 'sql',
        'html', 'css', 'programming', 'software', 'developer',
        'engineer', 'data', 'analysis', 'machine learning', 'ai',
        
        # Time indicators
        'january', 'february', 'march', 'april', 'may', 'june',
        'july', 'august', 'september', 'october', 'november', 'december',
        '2020', '2021', '2022', '2023', '2024', '2025',
        'present', 'current', 'ongoing',
        
        # Contact indicators
        'email', 'phone', 'linkedin', 'github', 'portfolio',
    ]
    
    text_lower = text.lower()
    found_indicators = sum(1 for indicator in resume_indicators if indicator in text_lower)
    
    if found_indicators < 2:
        return False, "File doesn't appear to be a resume. Please upload a valid resume."
    
    # Check if mostly gibberish
    alpha_chars = sum(c.isalpha() for c in text)
    total_chars = len(text.replace(' ', ''))
    
    if total_chars > 0 and alpha_chars / total_chars < 0.3:
        return False, "Extracted text appears corrupted."
    
    return True, ""
